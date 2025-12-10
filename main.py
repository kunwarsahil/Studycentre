import os
import json
import logging
import openai
from openai import OpenAI
import tempfile
from pathlib import Path
from typing import Optional, List, Dict, Tuple
from urllib.parse import urlparse
import requests
import pdfplumber
import docx
import eml_parser
import uuid
from dotenv import load_dotenv
from datetime import datetime, timedelta

# --- FastAPI Imports ---
from fastapi import FastAPI, UploadFile, File, HTTPException, Body
from pydantic import BaseModel, Field
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy import create_engine, Column, String, Integer, Float, DateTime, Text, Boolean
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, Session
from sqlalchemy.sql import func

# --- Load the .env file ---
load_dotenv()

# --- Setup Logging ---
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ######################################################################
# # 1. AGENT CLASS DEFINITIONS (All your code, unchanged)
# ######################################################################

# (All 4 of your classes: DocumentTextExtractor, FlashcardGenerator, QuizAgent, DoubtAgent go here...)
# (I'm omitting them for brevity, but they are *required* in this file)
class DocumentTextExtractor:
    """
    Handles downloading and text extraction from PDF, DOCX, and EML files.
    """
    SUPPORTED_FORMATS = {'pdf', 'docx', 'eml'}
    MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB
    REQUEST_TIMEOUT = 30  # seconds

    def __init__(
        self,
        max_file_size: int = MAX_FILE_SIZE,
        request_timeout: int = REQUEST_TIMEOUT
    ):
        self.max_file_size = max_file_size
        self.request_timeout = request_timeout

    def _table_to_markdown(self, table_data: list) -> str:
        """Converts extracted table data into markdown format."""
        if not table_data or not any(table_data):
            return ""
        try:
            num_cols = max(len(row) for row in table_data if row)
            markdown_lines = []
            
            header = [(cell if cell is not None else "") for cell in table_data[0]]
            header += [""] * (num_cols - len(header))
            markdown_lines.append(
                "| " + " | ".join(str(cell).replace("|", "\\|") for cell in header) + " |"
            )
            markdown_lines.append("|" + "|".join(["---"] * num_cols) + "|")
            
            for row in table_data[1:]:
                if not row: continue
                row = [(cell if cell is not None else "") for cell in row]
                row += [""] * (num_cols - len(row))
                markdown_lines.append(
                    "| " + " | ".join(str(cell).replace("|", "\\|") for cell in row) + " |"
                )
            return "\n".join(markdown_lines)
        except Exception as e:
            logger.warning(f"Failed to convert table to markdown: {e}")
            return ""

    def _is_url(self, path: str) -> bool:
        """Check if a string is a URL."""
        try:
            result = urlparse(path)
            return all([result.scheme, result.netloc])
        except Exception:
            return False

    def _download_file(self, url: str) -> Path:
        """Download file from URL to temporary location with safety checks."""
        try:
            parsed = urlparse(url)
            if parsed.scheme not in ('http', 'https'):
                raise ValueError(f"Unsupported URL scheme: {parsed.scheme}")
            
            response = requests.get(
                url,
                timeout=self.request_timeout,
                stream=True,
                allow_redirects=True
            )
            response.raise_for_status()
            
            content_length = response.headers.get('content-length')
            if content_length and int(content_length) > self.max_file_size:
                raise ValueError(f"File too large: {content_length} bytes")
            
            file_extension = Path(parsed.path).suffix.lower().lstrip('.')
            if not file_extension:
                content_type = response.headers.get('content-type', '')
                ext_map = {
                    'application/pdf': 'pdf',
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
                    'message/rfc822': 'eml'
                }
                file_extension = ext_map.get(content_type.split(';')[0], 'bin')
            
            temp_file = tempfile.NamedTemporaryFile(
                delete=False,
                suffix=f".{file_extension}"
            )
            
            downloaded_size = 0
            for chunk in response.iter_content(chunk_size=8192):
                downloaded_size += len(chunk)
                if downloaded_size > self.max_file_size:
                    temp_file.close()
                    os.unlink(temp_file.name)
                    raise ValueError(f"File exceeds maximum size: {self.max_file_size}")
                temp_file.write(chunk)
            
            temp_file.close()
            logger.info(f"Downloaded {downloaded_size} bytes to {temp_file.name}")
            return Path(temp_file.name)
            
        except requests.RequestException as e:
            logger.error(f"Failed to download {url}: {e}")
            raise
        except Exception as e:
            logger.error(f"Unexpected error downloading {url}: {e}")
            raise

    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        full_text = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        text = page.extract_text()
                        if text:
                            full_text.append(text)
                        for table in page.extract_tables():
                            if table:
                                markdown = self._table_to_markdown(table)
                                if markdown:
                                    full_text.append(f"\n{markdown}\n")
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num}: {e}")
                        continue
        except Exception as e:
            logger.error(f"Failed to open PDF {file_path}: {e}")
            raise
        return "\n".join(full_text)

    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        full_text = []
        try:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            for table in doc.tables:
                table_data = [[cell.text for cell in row.cells] for row in table.rows]
                markdown = self._table_to_markdown(table_data)
                if markdown:
                    full_text.append(f"\n{markdown}\n")
        except Exception as e:
            logger.error(f"Failed to extract DOCX {file_path}: {e}")
            raise
        return "\n".join(full_text)

    def _extract_eml(self, file_path: Path) -> str:
        """Extract text from EML file."""
        try:
            with open(file_path, 'rb') as f:
                raw_email = f.read()
            ep = eml_parser.EmlParser()
            parsed_eml = ep.decode_email_bytes(raw_email)
            text_parts = []
            for body_part in parsed_eml.get('body', []):
                if body_part.get('content_type') == 'text/plain':
                    text_parts.append(body_part.get('content', ''))
            if not text_parts and parsed_eml.get('body'):
                 for body_part in parsed_eml.get('body', []):
                    content = body_part.get('content', '')
                    if content:
                        text_parts.append(content)
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Failed to extract EML {file_path}: {e}")
            raise

    def extract_text(self, source: str) -> str:
        """Main public method to extract text from a URL or local file."""
        temp_file = None
        try:
            if self._is_url(source):
                logger.info(f"Downloading document from {source}")
                local_path = self._download_file(source)
                temp_file = local_path
            else:
                local_path = Path(source)
                if not local_path.exists():
                    raise FileNotFoundError(f"File not found: {source}")
            
            file_extension = local_path.suffix.lower().lstrip('.')
            if file_extension not in self.SUPPORTED_FORMATS:
                raise ValueError(
                    f"Unsupported file type: {file_extension}. "
                    f"Supported: {', '.join(self.SUPPORTED_FORMATS)}"
                )
            
            logger.info(f"Extracting text from {file_extension.upper()} file")
            if file_extension == 'pdf':
                full_text = self._extract_pdf(local_path)
            elif file_extension == 'docx':
                full_text = self._extract_docx(local_path)
            elif file_extension == 'eml':
                full_text = self._extract_eml(local_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            return full_text.strip()
            
        finally:
            if temp_file and temp_file.exists():
                try:
                    temp_file.unlink()
                    logger.debug(f"Cleaned up temporary file: {temp_file}")
                except Exception as e:
                    logger.warning(f"Failed to delete temporary file {temp_file}: {e}")

class FlashcardGenerator:
    """
    Uses the OpenAI API to generate flashcards from a given text.
    """
    def __init__(self):
        try:
            self.client = OpenAI()
        except openai.OpenAIError as e:
            logger.error(f"Failed to initialize OpenAI client. Is OPENAI_API_KEY set? Error: {e}")
            raise

    def _create_prompt(self, text: str) -> (str, str):
        """Creates the system and user prompt for the API call."""
        system_prompt = """
You are an expert study assistant. Your task is to generate flashcards
from the provided text. Analyze the text and identify key concepts,
definitions, and important facts.
Return your answer as a single, valid JSON object. This object should
contain one key: "flashcards". The value of "flashcards" should be a
list of JSON objects, where each object has two keys: "q" (for the question)
and "a" (for the answer).
Do not provide any preamble, explanation, or any text outside of the
single JSON object.
"""
        user_prompt = f"""
Here is the text to analyze:
--- TEXT START ---
{text}
--- TEXT END ---
Please generate at least 5-10 high-quality flashcards based on this text,
following the JSON format described in your instructions.
"""
        return system_prompt, user_prompt

    def generate(self, text: str) -> List[Dict[str, str]]:
        """
        Generates flashcards from the text.
        """
        if not text or not text.strip():
            logger.warning("No text provided to generate flashcards.")
            return []
            
        system_prompt, user_prompt = self._create_prompt(text)
        
        try:
            logger.info("Sending request to OpenAI API to generate flashcards...")
            completion = self.client.chat.completions.create(
                model="gpt-4o-mini", 
                response_format={"type": "json_object"}, 
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ]
            )
            response_content = completion.choices[0].message.content
            data = json.loads(response_content)
            flashcards = data.get("flashcards", [])
            
            if not flashcards:
                logger.warning("OpenAI response was valid JSON but contained no flashcards.")
                return []
                
            logger.info(f"Successfully generated {len(flashcards)} flashcards.")
            return flashcards

        except json.JSONDecodeError as e:
            logger.error(f"Failed to decode JSON from OpenAI response: {e}")
            logger.error(f"Raw response: {response_content}")
            return []
        except openai.OpenAIError as e:
            logger.error(f"An error occurred with the OpenAI API: {e}")
            return []
        except Exception as e:
            logger.error(f"An unexpected error occurred during flashcard generation: {e}")
            return []

class QuizAgent:
    """
    Generates quizzes and grades them in a single batch.
    Tracks accuracy and adjusts difficulty.
    """
    def __init__(self):
        try:
            self.client = OpenAI()
            self.accuracy_tracker = []
        except openai.OpenAIError as e:
            logger.error(f"Failed to initialize OpenAI client. Is OPENAI_API_KEY set? Error: {e}")
            raise

    def generate_quiz(self, context: str, difficulty: str = "Easy", num_questions: int = 5) -> List[Dict[str, str]]:
        """Generates a new quiz from the context text. (1 API Call)"""
        logger.info(f"Generating a {difficulty} quiz with {num_questions} questions...")
        
        system_prompt = f"""
You are an expert quiz designer. Your task is to create a {num_questions}-question
quiz based on the provided text. The difficulty must be {difficulty}.
- **Easy** questions should be simple definitions or "what is" questions.
- **Medium** questions should ask to explain a concept or a "how to" process.
- **Hard** questions should ask for analysis, comparison, or "why" something works.
Return your answer as a single, valid JSON object. This object should
contain one key: "questions". The value of "questions" should be a
list of JSON objects, where each object has two keys: "q" (for the question)
and "a" (for the correct answer).
Do not provide any preamble or text outside of the single JSON object.
"""
        user_prompt = f"""
Here is the text to analyze:
--- TEXT START ---
{context}
--- TEXT END ---
"""
        try:
            completion = self.client.chat.completions.create(
                model="gpt-4o-mini", 
                response_format={"type": "json_object"}, 
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ]
            )
            data = json.loads(completion.choices[0].message.content)
            quiz = data.get("questions", [])
            if not quiz:
                logger.warning("OpenAI response was valid JSON but contained no questions.")
                return []
            logger.info(f"Successfully generated {len(quiz)} questions.")
            return quiz
        except Exception as e:
            logger.error(f"Failed to generate quiz: {e}")
            return []

    def grade_quiz(self, quiz_data: List[Dict[str, str]], user_answers: List[str]) -> List[Dict[str, any]]:
        """Grades an entire quiz in a single batch. (1 API Call)"""
        logger.info("Grading quiz in a single batch...")
        
        grading_data = []
        for i, item in enumerate(quiz_data):
            grading_data.append({
                "question": item['q'],
                "correct_answer": item['a'],
                "user_answer": user_answers[i] if i < len(user_answers) else ""
            })
            
        system_prompt = """
You are a strict but fair teaching assistant. Your job is to evaluate a
student's quiz answers. You will be given a JSON list of quiz items.
Evaluate each [Student's Answer] for correctness. The student does not
need to be word-for-word perfect, but they must capture the main idea.
Respond *only* with a valid JSON object. This object must have one
key: "results". The value should be a list of JSON objects, one for
each question, in the *same order*. Each object must have two keys:
1. "is_correct": a boolean (true or false).
2. "feedback": a short string (1-2 sentences) explaining why the answer
   is correct or incorrect.
"""
        user_prompt = f"""
Here is the quiz to grade:
{json.dumps(grading_data, indent=2)}
"""
        try:
            completion = self.client.chat.completions.create(
                model="gpt-4o-mini",
                response_format={"type": "json_object"},
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ]
            )
            results = json.loads(completion.choices[0].message.content)
            self.track_results(results.get("results", []))
            return results.get("results", [])
        except Exception as e:
            logger.error(f"Failed to grade quiz: {e}")
            return [{"is_correct": False, "feedback": "Error grading this question."}] * len(quiz_data)

    def track_results(self, results: List[Dict[str, any]]):
        """Logs a batch of results to the accuracy tracker."""
        for item in results:
            self.accuracy_tracker.append(item.get("is_correct", False))
        logger.info(f"Logged {len(results)} new results to tracker.")

    def get_next_difficulty(self, current_difficulty: str) -> str:
        """Analyzes the accuracy tracker and suggests the next difficulty."""
        if not self.accuracy_tracker:
            logger.info("No accuracy data, starting fresh.")
            return "Easy"
            
        score = sum(self.accuracy_tracker) / len(self.accuracy_tracker)
        self.accuracy_tracker.clear()
        logger.info(f"Quiz complete. Score: {score*100:.0f}%")
        
        if score > 0.8:
            if current_difficulty == "Easy":
                next_level = "Medium"
            else:
                next_level = "Hard"
            logger.info("Great job! Increasing difficulty.")
        elif score < 0.5:
            if current_difficulty == "Hard":
                next_level = "Medium"
            else:
                next_level = "Easy"
            logger.info("Let's review. Decreasing difficulty.")
        else:
            next_level = current_difficulty
            logger.info("Good work. Staying at this difficulty.")
            
        return next_level

class DoubtAgent:
    """
    Answers contextual questions based *only* on the provided text.
    """
    def __init__(self):
        try:
            self.client = OpenAI()
        except openai.OpenAIError as e:
            logger.error(f"Failed to initialize OpenAI client. Is OPENAI_API_KEY set? Error: {e}")
            raise

    def ask_question(self, question: str, context: str) -> Dict[str, str]:
        """Asks a question and gets a contextual answer."""
        if not context or not context.strip():
            logger.warning("No context provided.")
            return {"answer": "Error: No document context was provided.", "reference": ""}
        if not question:
            return {"answer": "Error: No question was asked.", "reference": ""}

        logger.info(f"Asking contextual question: {question[:50]}...")
        
        system_prompt = """
You are an expert tutor and study assistant. Your job is to answer the
student's question based *only* on the provided [Context].

**Your Rules:**
1.  Read the [Question] and find the answer within the [Context].
2.  Provide a clear, helpful explanation. If the context provides an
    example, use it.
3.  **Crucially:** You MUST NOT use any outside knowledge.
4.  If the answer is not in the [Context], you must state: "I'm sorry,
    that information is not available in the provided document."
5.  After your answer, provide the exact quote(s) from the context you
    used as a reference.

**Respond *only* with a valid JSON object** with two keys:
1.  "answer": (Your clear explanation)
2.  "reference": (The exact quote(s) from the context, or an empty
    string if the answer wasn't found)
"""
        user_prompt = f"""
[Context]:
{context}

---
[Question]:
{question}
"""
        try:
            completion = self.client.chat.completions.create(
                model="gpt-4o-mini",
                response_format={"type": "json_object"},
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ]
            )
            data = json.loads(completion.choices[0].message.content)
            return data
        except json.JSONDecodeError as e:
            logger.error(f"Failed to decode JSON from OpenAI: {e}")
            return {"answer": "Error: Failed to get a valid response from AI.", "reference": ""}
        except openai.OpenAIError as e:
            logger.error(f"OpenAI API error: {e}")
            return {"answer": f"Error: {e}", "reference": ""}

class PlannerAgent:
    """
    Builds a smart revision schedule based on topic weightage and user progress.
    """
    def __init__(self):
        try:
            self.client = OpenAI()
        except openai.OpenAIError as e:
            logger.error(f"Failed to initialize OpenAI client. Is OPENAI_API_KEY set? Error: {e}")
            raise

    def create_revision_plan(
        self, 
        context: str, 
        exam_date: Optional[str] = None,
        topics: Optional[List[str]] = None,
        current_performance: Optional[Dict[str, float]] = None
    ) -> Dict:
        """Creates a revision schedule based on topics and performance."""
        logger.info("Creating revision plan...")
        
        # Default exam date is 30 days from now
        if exam_date:
            try:
                exam_dt = datetime.fromisoformat(exam_date.replace('Z', '+00:00'))
            except:
                exam_dt = datetime.now() + timedelta(days=30)
        else:
            exam_dt = datetime.now() + timedelta(days=30)
        
        days_until_exam = (exam_dt - datetime.now()).days
        days_until_exam = max(1, days_until_exam)
        
        system_prompt = """
You are an expert study planner. Your task is to create a revision schedule
based on the provided study material and student performance data.

Analyze the topics in the material and create a daily revision plan that:
1. Allocates more time to difficult topics (low performance scores)
2. Balances revision across all topics
3. Includes spaced repetition (revisit topics at increasing intervals)
4. Allows for breaks and review days

Return a JSON object with:
- "plan": A list of daily plans, where each day has:
  - "day": Day number (1, 2, 3...)
  - "date": ISO format date string
  - "topics": List of topics to review
  - "duration_minutes": Estimated study time
  - "focus": "high" | "medium" | "low" priority
- "summary": A brief explanation of the plan strategy
"""
        user_prompt = f"""
Study Material Context:
{context[:5000]}

Days until exam: {days_until_exam}
Topics to cover: {topics if topics else "All topics from the material"}
Performance data: {json.dumps(current_performance) if current_performance else "No performance data available"}

Create a revision plan that optimizes learning and retention.
"""
        try:
            completion = self.client.chat.completions.create(
                model="gpt-4o-mini",
                response_format={"type": "json_object"},
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ]
            )
            data = json.loads(completion.choices[0].message.content)
            
            # Enhance plan with actual dates
            plan = data.get("plan", [])
            start_date = datetime.now()
            for i, day_plan in enumerate(plan):
                day_plan["date"] = (start_date + timedelta(days=i)).isoformat()
            
            return {
                "plan": plan,
                "summary": data.get("summary", ""),
                "exam_date": exam_dt.isoformat(),
                "days_until_exam": days_until_exam
            }
        except Exception as e:
            logger.error(f"Failed to create revision plan: {e}")
            return {
                "plan": [],
                "summary": f"Error creating plan: {str(e)}",
                "exam_date": exam_dt.isoformat(),
                "days_until_exam": days_until_exam
            }

    def get_topic_analysis(self, context: str) -> List[Dict[str, str]]:
        """Analyzes the document and extracts topics with their importance."""
        logger.info("Analyzing topics from document...")
        
        system_prompt = """
You are a study material analyzer. Extract all major topics and subtopics
from the provided text, and assign each a weight (importance) from 1-10.

Return a JSON object with:
- "topics": A list of objects, each with:
  - "name": Topic name
  - "weight": Importance score (1-10)
  - "description": Brief description of the topic
"""
        user_prompt = f"""
Analyze this study material and extract topics:
{context[:5000]}
"""
        try:
            completion = self.client.chat.completions.create(
                model="gpt-4o-mini",
                response_format={"type": "json_object"},
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ]
            )
            data = json.loads(completion.choices[0].message.content)
            return data.get("topics", [])
        except Exception as e:
            logger.error(f"Failed to analyze topics: {e}")
            return []


# ######################################################################
# # 2. DATABASE SETUP
# ######################################################################

Base = declarative_base()

class Document(Base):
    __tablename__ = "documents"
    id = Column(String, primary_key=True)
    filename = Column(String)
    text_content = Column(Text)
    uploaded_at = Column(DateTime, default=datetime.utcnow)
    text_length = Column(Integer)

class Flashcard(Base):
    __tablename__ = "flashcards"
    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    document_id = Column(String)
    question = Column(Text)
    answer = Column(Text)
    created_at = Column(DateTime, default=datetime.utcnow)
    last_reviewed = Column(DateTime)
    review_count = Column(Integer, default=0)
    mastery_level = Column(Float, default=0.0)  # 0.0 to 1.0

class QuizResult(Base):
    __tablename__ = "quiz_results"
    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    document_id = Column(String)
    difficulty = Column(String)
    score = Column(Float)  # 0.0 to 1.0
    total_questions = Column(Integer)
    completed_at = Column(DateTime, default=datetime.utcnow)

class PerformanceMetric(Base):
    __tablename__ = "performance_metrics"
    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    document_id = Column(String)
    topic = Column(String)
    score = Column(Float)  # 0.0 to 1.0
    metric_type = Column(String)  # 'quiz', 'flashcard', etc.
    recorded_at = Column(DateTime, default=datetime.utcnow)

class RevisionPlan(Base):
    __tablename__ = "revision_plans"
    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    document_id = Column(String)
    plan_data = Column(Text)  # JSON string
    exam_date = Column(DateTime)
    created_at = Column(DateTime, default=datetime.utcnow)

# Initialize database
DATABASE_URL = "sqlite:///./study_assistant.db"
engine = create_engine(DATABASE_URL, connect_args={"check_same_thread": False})
Base.metadata.create_all(bind=engine)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)


# ######################################################################
# # 3. FASTAPI SETUP
# ######################################################################

app = FastAPI(title="Study Agent API")

# --- ADD THIS: CORS MIDDLEWARE ---
# This is required to allow your HTML frontend to call the backend.
origins = [
    "*",  # This allows all origins (e.g., file://, localhost)
    # "http://localhost",
    # "http://localhost:8080",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],  # Allows all methods (GET, POST, etc.)
    allow_headers=["*"],
)
# --- END OF CORS SECTION ---


# --- Global Agent Instances ---
extractor = DocumentTextExtractor()
flashcard_gen = FlashcardGenerator()
quiz_agent = QuizAgent()
doubt_agent = DoubtAgent()
planner_agent = PlannerAgent()

# --- Database Session Dependency ---
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

MAX_TEXT_FOR_AI = 8000


# ######################################################################
# # 4. PYDANTIC DATA MODELS
# ######################################################################

# --- Base Models ---
class QAPair(BaseModel):
    q: str
    a: str

class GradeResult(BaseModel):
    is_correct: bool
    feedback: str

# --- Request Bodies ---
class ContextIDRequest(BaseModel):
    context_id: str = Field(..., description="The unique ID of the uploaded document context.")

class GenerateQuizRequest(BaseModel):
    context_id: str
    difficulty: str = "Easy"
    num_questions: int = 5

class GradeQuizRequest(BaseModel):
    quiz_data: List[QAPair]
    user_answers: List[str]
    context_id: str
    difficulty: str = "Easy"

class AskDoubtRequest(BaseModel):
    context_id: str
    question: str

# --- Response Bodies ---
class UploadResponse(BaseModel):
    context_id: str
    filename: str
    text_length: int
    text_snippet: str

class FlashcardResponse(BaseModel):
    flashcards: List[QAPair]

class QuizResponse(BaseModel):
    questions: List[QAPair]

class GradeResponse(BaseModel):
    results: List[GradeResult]

class DoubtResponse(BaseModel):
    answer: str
    reference: str

class CreatePlanRequest(BaseModel):
    context_id: str
    exam_date: Optional[str] = None
    topics: Optional[List[str]] = None

class PlanResponse(BaseModel):
    plan: List[Dict]
    summary: str
    exam_date: str
    days_until_exam: int

class TopicAnalysisResponse(BaseModel):
    topics: List[Dict[str, str]]

class PerformanceStatsResponse(BaseModel):
    document_id: str
    total_quizzes: int
    average_score: float
    total_flashcards: int
    mastered_flashcards: int
    topic_performance: Dict[str, float]


# ######################################################################
# # 5. API ENDPOINTS
# ######################################################################

def get_context_text(context_id: str, db: Session) -> str:
    """Helper function to retrieve and validate context text from database."""
    doc = db.query(Document).filter(Document.id == context_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Context ID not found. Please upload the document first.")
    return doc.text_content

@app.get("/", summary="Health Check")
def read_root():
    """A simple 'hello world' endpoint to check if the server is running."""
    return {"message": "Study Agent API is running."}

@app.post("/upload", response_model=UploadResponse, summary="Upload a Document")
def upload_document(file: UploadFile = File(...)):
    """
    Upload a .pdf, .docx, or .eml file.
    Extracts the text and returns a `context_id` for use with other endpoints.
    """
    db_session = next(get_db())
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=file.filename) as temp_f:
            temp_f.write(file.file.read())
            temp_file_path = temp_f.name
        
        text = extractor.extract_text(temp_file_path)
        os.unlink(temp_file_path)
        
        if not text:
            raise HTTPException(status_code=400, detail="Could not extract any text from the document.")

        context_id = str(uuid.uuid4())
        
        # Save to database
        doc = Document(
            id=context_id,
            filename=file.filename,
            text_content=text,
            text_length=len(text)
        )
        db_session.add(doc)
        db_session.commit()
        
        return UploadResponse(
            context_id=context_id,
            filename=file.filename,
            text_length=len(text),
            text_snippet=text[:200] + "..."
        )

    except Exception as e:
        logger.error(f"Error during file upload: {e}")
        db_session.rollback()
        raise HTTPException(status_code=500, detail=f"An error occurred: {str(e)}")
    finally:
        db_session.close()

@app.post("/flashcards", response_model=FlashcardResponse, summary="Generate Flashcards")
def get_flashcards(request: ContextIDRequest):
    """
    Generates flashcards from an uploaded document's text.
    """
    db = next(get_db())
    try:
        text = get_context_text(request.context_id, db)
        flashcards = flashcard_gen.generate(text[:MAX_TEXT_FOR_AI])
        
        # Save flashcards to database
        for fc in flashcards:
            flashcard = Flashcard(
                document_id=request.context_id,
                question=fc["q"],
                answer=fc["a"]
            )
            db.add(flashcard)
        db.commit()
        
        return FlashcardResponse(flashcards=flashcards)
    finally:
        db.close()

@app.post("/quiz/generate", response_model=QuizResponse, summary="Generate a Quiz")
def generate_quiz(request: GenerateQuizRequest):
    """
    Generates a quiz with a specific difficulty from the document's text.
    """
    db = next(get_db())
    try:
        text = get_context_text(request.context_id, db)
        quiz = quiz_agent.generate_quiz(
            context=text[:MAX_TEXT_FOR_AI],
            difficulty=request.difficulty,
            num_questions=request.num_questions
        )
        return QuizResponse(questions=quiz)
    finally:
        db.close()

@app.post("/quiz/grade", response_model=GradeResponse, summary="Grade a Quiz")
def grade_quiz(request: GradeQuizRequest):
    """
    Grades a completed quiz and saves results to database.
    """
    db = next(get_db())
    try:
        results = quiz_agent.grade_quiz(
            quiz_data=[item.model_dump() for item in request.quiz_data],
            user_answers=request.user_answers
        )
        
        # Calculate score
        correct_count = sum(1 for r in results if r.get("is_correct", False))
        score = correct_count / len(results) if results else 0.0
        
        # Save quiz result
        quiz_result = QuizResult(
            document_id=request.context_id,
            difficulty=request.difficulty,
            score=score,
            total_questions=len(results)
        )
        db.add(quiz_result)
        db.commit()
        
        return GradeResponse(results=results)
    finally:
        db.close()

@app.post("/doubt", response_model=DoubtResponse, summary="Ask a Contextual Question")
def ask_doubt(request: AskDoubtRequest):
    """
    Asks a "doubt" question about the document's content.
    """
    db = next(get_db())
    try:
        text = get_context_text(request.context_id, db)
        answer = doubt_agent.ask_question(
            question=request.question,
            context=text 
        )
        return DoubtResponse(answer=answer.get("answer", ""), reference=answer.get("reference", ""))
    finally:
        db.close()

@app.post("/planner/create", response_model=PlanResponse, summary="Create Revision Plan")
def create_revision_plan(request: CreatePlanRequest):
    """
    Creates a revision plan based on the document and optional exam date.
    """
    db = next(get_db())
    try:
        text = get_context_text(request.context_id, db)
        
        # Get performance data
        performance_metrics = db.query(PerformanceMetric).filter(
            PerformanceMetric.document_id == request.context_id
        ).all()
        performance_dict = {}
        for metric in performance_metrics:
            if metric.topic not in performance_dict:
                performance_dict[metric.topic] = []
            performance_dict[metric.topic].append(metric.score)
        # Average performance per topic
        performance_dict = {k: sum(v)/len(v) if v else 0.5 for k, v in performance_dict.items()}
        
        plan_data = planner_agent.create_revision_plan(
            context=text[:MAX_TEXT_FOR_AI],
            exam_date=request.exam_date,
            topics=request.topics,
            current_performance=performance_dict if performance_dict else None
        )
        
        # Save plan to database
        revision_plan = RevisionPlan(
            document_id=request.context_id,
            plan_data=json.dumps(plan_data),
            exam_date=datetime.fromisoformat(plan_data["exam_date"].replace('Z', '+00:00'))
        )
        db.add(revision_plan)
        db.commit()
        
        return PlanResponse(**plan_data)
    finally:
        db.close()

@app.post("/planner/topics", response_model=TopicAnalysisResponse, summary="Analyze Topics")
def analyze_topics(request: ContextIDRequest):
    """
    Analyzes the document and extracts topics with their importance.
    """
    db = next(get_db())
    try:
        text = get_context_text(request.context_id, db)
        topics = planner_agent.get_topic_analysis(text[:MAX_TEXT_FOR_AI])
        return TopicAnalysisResponse(topics=topics)
    finally:
        db.close()

@app.get("/performance/{context_id}", response_model=PerformanceStatsResponse, summary="Get Performance Stats")
def get_performance_stats(context_id: str):
    """
    Gets performance statistics for a document.
    """
    db = next(get_db())
    try:
        # Quiz results
        quiz_results = db.query(QuizResult).filter(QuizResult.document_id == context_id).all()
        total_quizzes = len(quiz_results)
        avg_score = sum(q.score for q in quiz_results) / total_quizzes if total_quizzes > 0 else 0.0
        
        # Flashcard stats
        flashcards = db.query(Flashcard).filter(Flashcard.document_id == context_id).all()
        total_flashcards = len(flashcards)
        mastered_flashcards = len([f for f in flashcards if f.mastery_level >= 0.8])
        
        # Topic performance
        performance_metrics = db.query(PerformanceMetric).filter(
            PerformanceMetric.document_id == context_id
        ).all()
        topic_performance = {}
        for metric in performance_metrics:
            if metric.topic not in topic_performance:
                topic_performance[metric.topic] = []
            topic_performance[metric.topic].append(metric.score)
        topic_performance = {k: sum(v)/len(v) if v else 0.0 for k, v in topic_performance.items()}
        
        return PerformanceStatsResponse(
            document_id=context_id,
            total_quizzes=total_quizzes,
            average_score=avg_score,
            total_flashcards=total_flashcards,
            mastered_flashcards=mastered_flashcards,
            topic_performance=topic_performance
        )
    finally:
        db.close()

@app.get("/documents", summary="List All Documents")
def list_documents():
    """
    Lists all uploaded documents.
    """
    db = next(get_db())
    try:
        documents = db.query(Document).all()
        return [{
            "id": doc.id,
            "filename": doc.filename,
            "uploaded_at": doc.uploaded_at.isoformat() if doc.uploaded_at else None,
            "text_length": doc.text_length
        } for doc in documents]
    finally:
        db.close()